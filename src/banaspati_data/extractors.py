from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Iterator

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
_OCR_CACHE: dict[str, str] = {}


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _record(
    source: Path,
    root: Path,
    text: str,
    content_type: str,
    page: int | None = None,
    section: str | None = None,
    extra: dict | None = None,
) -> dict:
    relative = source.relative_to(root).as_posix()
    stable_key = f"{relative}|{page}|{section}|{content_type}|{text[:100]}"
    metadata = {
        "source_file": relative,
        "file_name": source.name,
        "file_type": source.suffix.lower().lstrip("."),
        "page": page,
        "section": section,
        "content_type": content_type,
    }
    if extra:
        metadata.update(extra)
    return {
        "id": hashlib.sha1(stable_key.encode()).hexdigest(),
        "text": _clean_text(text),
        "metadata": metadata,
    }


def _ocr_image(image_bytes: bytes) -> str:
    image_hash = hashlib.sha1(image_bytes).hexdigest()
    if image_hash in _OCR_CACHE:
        return _OCR_CACHE[image_hash]
    try:
        import pytesseract
        from PIL import Image, ImageEnhance
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    except ImportError:
        return ""
    try:
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        if max(image.size) < 2500:
            image = image.resize(
                (image.width * 3, image.height * 3), Image.Resampling.LANCZOS
            )
        image = ImageEnhance.Contrast(image).enhance(1.5)
        text = _clean_text(pytesseract.image_to_string(image, config="--psm 6"))
    except Exception:
        text = ""
    _OCR_CACHE[image_hash] = text
    return text


def _save_image(
    image_bytes: bytes,
    extension: str,
    assets_dir: Path | None,
    source: Path,
    root: Path,
    location: str,
) -> str | None:
    if assets_dir is None:
        return None
    source_key = source.relative_to(root).as_posix()
    name = hashlib.sha1(f"{source_key}|{location}".encode()).hexdigest()
    output = assets_dir / f"{name}.{extension.lstrip('.') or 'bin'}"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_bytes(image_bytes)
    return output.as_posix()


def extract_pdf(
    path: Path, root: Path, use_ocr: bool = False, assets_dir: Path | None = None
) -> Iterator[dict]:
    import pymupdf

    with pymupdf.open(path) as document:
        for page_number, page in enumerate(document, start=1):
            text = _clean_text(page.get_text("text"))
            if text:
                yield _record(path, root, text, "text", page=page_number)

            try:
                tables = page.find_tables().tables
            except Exception:
                tables = []
            for table_index, table in enumerate(tables, start=1):
                rows = table.extract()
                table_text = "\n".join(
                    " | ".join(_clean_text(str(cell or "")) for cell in row) for row in rows
                )
                if table_text.strip(" |\n"):
                    yield _record(
                        path,
                        root,
                        table_text,
                        "table",
                        page=page_number,
                        section=f"table-{table_index}",
                    )

            for image_index, image in enumerate(page.get_images(full=True), start=1):
                extracted = document.extract_image(image[0])
                image_bytes = extracted.get("image", b"")
                asset_path = _save_image(
                    image_bytes,
                    extracted.get("ext", "bin"),
                    assets_dir,
                    path,
                    root,
                    f"page-{page_number}-image-{image_index}",
                )
                ocr_text = _ocr_image(image_bytes) if use_ocr else ""
                description = ocr_text or f"[Gambar {image_index} pada halaman {page_number}]"
                yield _record(
                    path,
                    root,
                    description,
                    "image_ocr" if ocr_text else "image",
                    page=page_number,
                    section=f"image-{image_index}",
                    extra={"has_ocr_text": bool(ocr_text), "asset_path": asset_path},
                )


def extract_docx(
    path: Path, root: Path, use_ocr: bool = False, assets_dir: Path | None = None
) -> Iterator[dict]:
    from docx import Document

    document = Document(path)
    section_number = 1
    buffer: list[str] = []

    def flush() -> dict | None:
        nonlocal buffer
        text = _clean_text("\n".join(buffer))
        buffer = []
        return _record(path, root, text, "text", section=f"section-{section_number}") if text else None

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading") and buffer:
            item = flush()
            if item:
                yield item
            section_number += 1
        buffer.append(text)
    item = flush()
    if item:
        yield item

    for table_index, table in enumerate(document.tables, start=1):
        table_text = "\n".join(
            " | ".join(_clean_text(cell.text) for cell in row.cells) for row in table.rows
        )
        if table_text.strip(" |\n"):
            yield _record(path, root, table_text, "table", section=f"table-{table_index}")

    for image_index, relation in enumerate(document.part.rels.values(), start=1):
        if "image" not in relation.reltype:
            continue
        image_bytes = relation.target_part.blob
        extension = Path(str(relation.target_part.partname)).suffix or ".bin"
        asset_path = _save_image(
            image_bytes, extension, assets_dir, path, root, f"image-{image_index}"
        )
        ocr_text = _ocr_image(image_bytes) if use_ocr else ""
        description = ocr_text or f"[Gambar {image_index} pada dokumen]"
        yield _record(
            path,
            root,
            description,
            "image_ocr" if ocr_text else "image",
            section=f"image-{image_index}",
            extra={"has_ocr_text": bool(ocr_text), "asset_path": asset_path},
        )


def extract_plain_text(path: Path, root: Path) -> Iterator[dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if _clean_text(text):
        yield _record(path, root, text, "text")


def extract_directory(
    input_dir: Path, use_ocr: bool = False, assets_dir: Path | None = None
) -> list[dict]:
    input_dir = input_dir.resolve()
    if not input_dir.exists():
        raise FileNotFoundError(f"Folder input tidak ditemukan: {input_dir}")
    if not input_dir.is_dir():
        raise NotADirectoryError(f"Path input bukan folder: {input_dir}")

    supported_files = [
        path
        for path in sorted(input_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    if not supported_files:
        formats = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(
            f"Tidak ada dokumen yang didukung di {input_dir}. Format yang didukung: {formats}"
        )

    records: list[dict] = []
    for path in supported_files:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            records.extend(extract_pdf(path, input_dir, use_ocr, assets_dir))
        elif suffix == ".docx":
            records.extend(extract_docx(path, input_dir, use_ocr, assets_dir))
        else:
            records.extend(extract_plain_text(path, input_dir))
    if not records:
        raise ValueError(
            "Dokumen ditemukan, tetapi tidak ada konten yang berhasil diekstrak. "
            "Untuk PDF hasil scan, coba jalankan kembali dengan --ocr."
        )
    return records
