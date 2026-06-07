from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from banaspati_data.chunking import chunk_documents
from banaspati_data.extractors import extract_directory
from banaspati_data.io_utils import read_jsonl, write_jsonl
from banaspati_data.vector_store import build_chroma


def test_missing_input_directory_fails_clearly(tmp_path: Path):
    import pytest

    with pytest.raises(FileNotFoundError, match="Folder input tidak ditemukan"):
        extract_directory(tmp_path / "missing")


def test_empty_directory_fails_clearly(tmp_path: Path):
    import pytest

    with pytest.raises(ValueError, match="Tidak ada dokumen yang didukung"):
        extract_directory(tmp_path)


def test_plain_text_extraction_keeps_source_metadata(tmp_path: Path):
    source = tmp_path / "aturan.md"
    source.write_text("# Aturan\nMahasiswa wajib hadir.", encoding="utf-8")

    records = extract_directory(tmp_path)

    assert len(records) == 1
    assert records[0]["metadata"]["source_file"] == "aturan.md"
    assert records[0]["metadata"]["content_type"] == "text"
    assert "wajib hadir" in records[0]["text"]


def test_recursive_chunking_keeps_traceability():
    documents = [
        {
            "id": "parent",
            "text": "Kalimat pertama. Kalimat kedua yang lebih panjang. Kalimat ketiga.",
            "metadata": {"source_file": "uji.txt", "page": 1, "content_type": "text"},
        }
    ]

    chunks = chunk_documents(documents, strategy="recursive", chunk_size=45, overlap=10)

    assert len(chunks) >= 2
    assert all(chunk["metadata"]["parent_id"] == "parent" for chunk in chunks)
    assert all(chunk["metadata"]["source_file"] == "uji.txt" for chunk in chunks)
    assert all(len(chunk["text"]) <= 45 for chunk in chunks)


def test_uncaptioned_image_placeholder_is_not_chunked():
    documents = [
        {
            "id": "image",
            "text": "[Gambar 1 pada halaman 2]",
            "metadata": {"source_file": "uji.pdf", "page": 2, "content_type": "image"},
        }
    ]

    assert chunk_documents(documents) == []


def test_jsonl_round_trip(tmp_path: Path):
    records = [{"id": "1", "text": "konten", "metadata": {"page": 2}}]
    output = tmp_path / "records.jsonl"

    assert write_jsonl(output, records) == 1
    assert read_jsonl(output) == records


def test_empty_chunks_are_rejected_before_loading_model(tmp_path: Path):
    import pytest

    with pytest.raises(ValueError, match="File chunk kosong"):
        build_chroma([], tmp_path / "chroma", "banaspati", "unused")


def test_docx_extraction(tmp_path: Path):
    from docx import Document

    docx = Document()
    docx.add_heading("Kurikulum", level=1)
    docx.add_paragraph("Total SKS wajib adalah 144.")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Kode"
    table.cell(0, 1).text = "Nama Mata Kuliah"
    docx.save(tmp_path / "kurikulum.docx")

    records = extract_directory(tmp_path, assets_dir=tmp_path / "assets")
    types = {(item["metadata"]["file_type"], item["metadata"]["content_type"]) for item in records}

    assert ("docx", "text") in types
    assert ("docx", "table") in types


def test_pdf_extraction(tmp_path: Path):
    import pytest

    pymupdf = pytest.importorskip("pymupdf")
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Jadwal akademik semester")
    pdf.save(tmp_path / "jadwal.pdf")
    pdf.close()

    records = extract_directory(tmp_path)

    assert any(item["metadata"]["file_type"] == "pdf" for item in records)
    assert any(item["metadata"]["page"] == 1 for item in records if item["metadata"]["file_type"] == "pdf")
