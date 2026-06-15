"""
Script untuk mengekstrak ulang Jadwal Perkuliahan.docx dengan metode khusus yang
mampu menangani tabel kompleks dan menambahkannya ke ChromaDB yang sudah ada.
"""
import sys
sys.path.insert(0, '.')

from docx import Document
from pathlib import Path
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document as LCDocument
import warnings
warnings.filterwarnings('ignore')

JADWAL_PATH = Path("data/database/Jadwal Perkuliahan.docx")

def extract_jadwal_tables(path: Path):
    """
    Ekstrak tabel jadwal dari .docx dengan presisi tinggi.
    Setiap baris tabel menjadi satu entri teks untuk retrieval.
    """
    doc = Document(path)
    records = []
    
    for table_idx, table in enumerate(doc.tables, start=1):
        # Ekstrak header tabel (baris pertama)
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        header_str = " | ".join(h for h in headers if h)
        
        # Ekstrak setiap baris data
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            # Lewati baris yang kosong
            if not any(cells):
                continue
            
            # Buat teks kontekstual: "Kolom: Nilai, Kolom: Nilai, ..."
            row_parts = []
            for h, c in zip(headers, cells):
                if c:
                    if h:
                        row_parts.append(f"{h}: {c}")
                    else:
                        row_parts.append(c)
            
            row_text = " | ".join(row_parts)
            
            # Juga buat versi gabungan seluruh baris satu tabel untuk konteks lebih luas
            records.append({
                "text": row_text,
                "metadata": {
                    "source_file": "Jadwal Perkuliahan.docx",
                    "file_name": "Jadwal Perkuliahan.docx",
                    "file_type": "docx",
                    "content_type": "table",
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "table_header": header_str
                }
            })
    
    # Juga ekstrak seluruh tabel sebagai blok besar untuk retrieval konteks penuh
    for table_idx, table in enumerate(doc.tables, start=1):
        all_rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                all_rows.append(row_text)
        
        if all_rows:
            full_table_text = "\n".join(all_rows)
            records.append({
                "text": full_table_text,
                "metadata": {
                    "source_file": "Jadwal Perkuliahan.docx",
                    "file_name": "Jadwal Perkuliahan.docx",
                    "file_type": "docx",
                    "content_type": "full_table",
                    "table_index": table_idx,
                }
            })
    
    # Juga ekstrak paragraf teks biasa
    for para_idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text and len(text) > 10:
            records.append({
                "text": text,
                "metadata": {
                    "source_file": "Jadwal Perkuliahan.docx",
                    "file_name": "Jadwal Perkuliahan.docx",
                    "file_type": "docx",
                    "content_type": "paragraph",
                    "para_index": para_idx
                }
            })
    
    return records


def main():
    print(f"Mengekstrak data dari: {JADWAL_PATH}")
    records = extract_jadwal_tables(JADWAL_PATH)
    print(f"Total records diekstrak: {len(records)}")
    
    # Tampilkan sample
    print("\n=== SAMPLE 3 RECORDS PERTAMA ===")
    for r in records[:3]:
        print(f"\n[{r['metadata']['content_type']}]")
        print(r['text'][:300])
    
    print("\n\nMemuat embedding model...")
    embed_model = HuggingFaceEmbeddings(model_name='paraphrase-multilingual-MiniLM-L12-v2')
    
    print("Menghubungkan ke ChromaDB...")
    vectorstore = Chroma(
        persist_directory='artifacts/chroma',
        collection_name='banaspati',
        embedding_function=embed_model
    )
    
    # Hapus chunks Jadwal Perkuliahan lama yang rusak
    print("Menghapus chunks Jadwal Perkuliahan lama yang rusak...")
    try:
        old_results = vectorstore.get(where={"source_file": "Jadwal Perkuliahan.docx"})
        if old_results and old_results['ids']:
            vectorstore.delete(ids=old_results['ids'])
            print(f"  Dihapus {len(old_results['ids'])} chunks lama.")
        else:
            print("  Tidak ada chunks lama ditemukan.")
    except Exception as e:
        print(f"  Warning saat menghapus: {e}")
    
    # Tambahkan records baru
    print(f"\nMenambahkan {len(records)} records baru ke ChromaDB...")
    docs_to_add = []
    for r in records:
        if r['text'].strip():
            docs_to_add.append(LCDocument(
                page_content=r['text'],
                metadata=r['metadata']
            ))
    
    # Tambahkan dalam batch
    batch_size = 100
    for i in range(0, len(docs_to_add), batch_size):
        batch = docs_to_add[i:i+batch_size]
        vectorstore.add_documents(batch)
        print(f"  Batch {i//batch_size + 1}: {len(batch)} docs ditambahkan.")
    
    print(f"\nSelesai! Total {len(docs_to_add)} chunks baru dari Jadwal Perkuliahan.docx berhasil diindeks.")
    
    # Verifikasi
    print("\n=== VERIFIKASI RETRIEVAL ===")
    test_docs = vectorstore.similarity_search("mata kuliah hari senin jadwal", k=5)
    jadwal_hits = [d for d in test_docs if 'Jadwal' in d.metadata.get('source_file', '')]
    print(f"Test query 'senin': {len(jadwal_hits)} dari 5 hasil berasal dari Jadwal Perkuliahan.")
    for d in jadwal_hits[:2]:
        print(f"\n  [{d.metadata.get('content_type')}] {d.page_content[:200]}")


if __name__ == "__main__":
    main()
